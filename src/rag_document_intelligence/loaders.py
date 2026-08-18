import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    DirectoryLoader,
    PyMuPDFLoader,
    TextLoader,
)

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from .config import DOCUMENTS_DIR

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load and normalize supported inputs into LangChain Documents."""

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".txt",
        ".md",
        ".docx",
    }

    def __init__(
        self,
        documents_dir: Path = DOCUMENTS_DIR,
    ):
        self.documents_dir = Path(documents_dir)

    def _validate_file(self, file_path: Path) -> None:
        """Validate that a file exists and has a supported extension."""
        if not file_path.exists():
            raise FileNotFoundError(
                f"Input file does not exist: {file_path}"
            )

        if not file_path.is_file():
            raise ValueError(
                f"Input path is not a file: {file_path}"
            )

        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            supported = ", ".join(
                sorted(self.SUPPORTED_EXTENSIONS)
            )
            raise ValueError(
                f"Unsupported file type: {file_path.suffix}. "
                f"Supported types: {supported}"
            )

        if file_path.stat().st_size == 0:
            raise ValueError(
                f"File is empty (0 bytes): {file_path.name}"
            )

    def _resolve_path(self, file_path: str | Path) -> Path:
        """Resolve a file path, preventing traversal for relative paths.

        Absolute paths are used as-is.  Relative paths that contain
        ``..`` are checked against the documents directory root.
        """
        path = Path(str(file_path))

        if not path.is_absolute() and ".." in path.parts:
            root = self.documents_dir.resolve()
            resolved = (root / path).resolve()
            try:
                resolved.relative_to(root)
            except ValueError:
                raise ValueError(
                    f"Path traversal detected; file must remain "
                    f"within {root}: {file_path}"
                )
            return resolved

        return path

    @staticmethod
    def _filter_empty(documents: List[Document]) -> List[Document]:
        """Remove documents whose content is empty after stripping."""
        filtered = [
            doc
            for doc in documents
            if doc.page_content.strip()
        ]
        if len(filtered) < len(documents):
            logger.debug(
                "Filtered %d empty document(s) from loader results",
                len(documents) - len(filtered),
            )
        return filtered

    def load_text_file(
        self,
        file_path: Path,
    ) -> List[Document]:
        """Load TXT or Markdown files."""
        loader = TextLoader(
            str(file_path),
            encoding="utf-8",
        )

        documents = loader.load()

        documents = self._filter_empty(documents)

        for document in documents:
            document.metadata["input_type"] = (
                file_path.suffix.lower().lstrip(".")
            )
            document.metadata["source_name"] = file_path.name

        return documents

    def load_pdf_file(
        self,
        file_path: Path,
    ) -> List[Document]:
        """Load a PDF document page by page."""
        loader = PyMuPDFLoader(str(file_path))
        documents = loader.load()

        documents = self._filter_empty(documents)

        for document in documents:
            document.metadata["input_type"] = "pdf"
            document.metadata["source_name"] = file_path.name

        return documents

    def load_docx_file(
        self,
        file_path: Path,
    ) -> List[Document]:
        """Load a DOCX document and preserve paragraph content."""
        if DocxDocument is None:
            raise RuntimeError(
                "python-docx is required to load DOCX files. "
                "Install it with: python -m pip install python-docx"
            )

        docx_document = DocxDocument(str(file_path))

        paragraphs = [
            paragraph.text.strip()
            for paragraph in docx_document.paragraphs
            if paragraph.text.strip()
        ]

        content = "\n".join(paragraphs)

        if not content:
            return []

        return [
            Document(
                page_content=content,
                metadata={
                    "source": str(file_path),
                    "source_name": file_path.name,
                    "input_type": "docx",
                },
            )
        ]

    def load_file(
        self,
        file_path: str | Path,
    ) -> List[Document]:
        """Load one supported file into normalized Documents."""
        path = self._resolve_path(file_path)
        self._validate_file(path)

        suffix = path.suffix.lower()

        if suffix == ".pdf":
            return self.load_pdf_file(path)

        if suffix in {".txt", ".md"}:
            return self.load_text_file(path)

        if suffix == ".docx":
            return self.load_docx_file(path)

        raise ValueError(
            f"Unsupported file type: {suffix}"
        )

    def load_text_content(
        self,
        text: str,
        source_name: str = "user_input",
    ) -> List[Document]:
        """Convert direct user-provided text into a Document."""
        if not text or not text.strip():
            raise ValueError(
                "text content cannot be empty"
            )

        return [
            Document(
                page_content=text.strip(),
                metadata={
                    "source": source_name,
                    "source_name": source_name,
                    "input_type": "text",
                },
            )
        ]

    def load_text_documents(self) -> List[Document]:
        """Load all TXT files from the document directory."""
        return self._load_directory(
            "*.txt"
        )

    def load_markdown_documents(self) -> List[Document]:
        """Load all Markdown files from the document directory."""
        return self._load_directory(
            "*.md"
        )

    def load_pdf_documents(self) -> List[Document]:
        """Load all PDF files recursively."""
        documents: List[Document] = []

        for path in self.documents_dir.rglob("*.pdf"):
            try:
                documents.extend(
                    self.load_pdf_file(path.resolve())
                )
            except Exception as exc:
                logger.warning(
                    "Failed to load PDF %s: %s", path, exc
                )

        return documents

    def load_docx_documents(self) -> List[Document]:
        """Load all DOCX files recursively."""
        documents: List[Document] = []

        for path in self.documents_dir.rglob("*.docx"):
            try:
                documents.extend(
                    self.load_docx_file(path.resolve())
                )
            except Exception as exc:
                logger.warning(
                    "Failed to load DOCX %s: %s", path, exc
                )

        return documents

    def _load_directory(
        self,
        pattern: str,
    ) -> List[Document]:
        """Load text-like files matching a pattern."""
        if not self.documents_dir.exists():
            logger.warning(
                "Documents directory does not exist: %s",
                self.documents_dir,
            )
            return []

        loader = DirectoryLoader(
            str(self.documents_dir),
            glob=pattern,
            loader_cls=TextLoader,
            loader_kwargs={
                "encoding": "utf-8"
            },
            show_progress=False,
        )

        documents = loader.load()

        extension = Path(pattern.replace("*", "")).suffix

        for document in documents:
            document.metadata["input_type"] = (
                extension.lstrip(".")
            )
            document.metadata["source_name"] = Path(
                document.metadata.get(
                    "source",
                    "unknown",
                )
            ).name

        return self._filter_empty(documents)

    def load_all_documents(self) -> List[Document]:
        """Load every supported document from the document directory."""
        documents: List[Document] = []

        if not self.documents_dir.exists():
            logger.warning(
                "Documents directory does not exist: %s",
                self.documents_dir,
            )
            return []

        for path in sorted(
            self.documents_dir.rglob("*")
        ):
            if not path.is_file():
                continue

            if path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
                continue

            try:
                documents.extend(
                    self.load_file(path.resolve())
                )
            except Exception as exc:
                logger.warning(
                    "Failed to load %s: %s", path, exc
                )

        return documents

    def supported_extensions(self) -> List[str]:
        """Return supported file extensions."""
        return sorted(self.SUPPORTED_EXTENSIONS)
